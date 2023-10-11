from docx import Document
from docx.shared import Inches

document = Document()

probe = 'AATAGTGAGACCGAAGATCGCGGCCGTAAGCCTCTGGGCACGGGT'
CpG = 'GC'
path = "results_20.fasta"

p = document.add_paragraph('')

with open(path, "r") as file:
    for line in file:
        p = document.add_paragraph('')
        splitted_line = line.split(CpG)
        print(splitted_line)
        for l in splitted_line:
            if l == CpG:
                p.add_run(l).bold = True
            else:
                p.add_run(l)
        document.add_page_break()


document.save('demo.docx')
